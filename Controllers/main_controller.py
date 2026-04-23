from PyQt6.QtWidgets import QMessageBox, QMainWindow

from Config.exceptions import InvalidMoveName, DuplicateMoveError
from Config.style import main_window_style
from Config.variables import Var

from Managers.builder import SetupBuilder
from Managers.display_manager import DisplayManager
from Managers.write_data_manager import WriteDataManager
from Managers.path_manager import PathManager

from docx import Document
from datetime import datetime

class MainController:
    def __init__(self):
        # =============== Build =============== #
        self.path_manager       = PathManager()
        self.display_manager    = DisplayManager()
        self.models             = SetupBuilder.build_models()
        self.views              = SetupBuilder.build_views()
        self.controllers        = SetupBuilder.build_controllers(self.models, self.views)

        # =============== Set Controllers Methods =============== #
        SetupBuilder.connect_controllers(self.controllers)
        self.controllers["move"].view.rename_move_method    = self.rename_move
        self.controllers["move"].global_remove_move         = self.global_remove_move
        self.views["navigator"].write_to_code_method        = self.write_to_code
        self.views["navigator"].show_view_method            = self.show_view
        self.views["navigator"].start_new_project_method    = self.start_new_project
        self.views["navigator"].close_project_method        = self.close_project

        # =============== Logic =============== #
        for name, controller in self.controllers.items():
            self.display_manager.register(name, controller)

        # =============== Root Widget =============== #
        self.root = QMainWindow()
        self.root.setWindowTitle("Config Tool")
        self.root.setCentralWidget(self.views["main"])
        self.root.setStyleSheet(main_window_style)
        self.root.show()
        self.root.showMaximized()  # show in full-screen

    def show_view(self, act):
        # Initialize app
        if act == "init":
            if self._initialize_app():
                Var.authority = self.views["navigator"].authority_combo.currentData()
                self.views["navigator"].authority_combo.setDisabled(True)
                self.display_manager.show("settings")
            return

        # Check if the project initialized
        if self.path_manager.path_project is None:
            QMessageBox.critical(self.views["main"], "שגיאה", "פרויקט לא מאותחל")
            return

        # Choose which view to show
        if act == "matrix":
            self.display_manager.show("matrix", self.models["move"].all_moves)
        elif act == "image":
            self.display_manager.show("image", self.controllers["move"].get_all_moves_names())
        elif act == "phue":
            self.display_manager.show("phue", self.models["image"].all_images, self.controllers["move"].get_all_moves_names())
        elif act == "parameters_ta":
            self.display_manager.show("parameters_ta", self.models["image"].all_images)
        elif act == "sk":
            self.display_manager.show("sk", self.models["move"].all_moves)
        else:
            self.display_manager.show(act)

    # ============================== CRUD ============================== #
    def rename_move(self, old_name, new_name):
        try:
            self.controllers["move"].rename_move(old_name, new_name)  # must be first
            self.controllers["matrix"].rename_move(old_name, new_name)
            self.controllers["sk"].rename_move(old_name, new_name)
            self.controllers["detector"].rename_move(old_name, new_name)
            self.controllers["image"].rename_move(old_name, new_name)
            self.controllers["phue"].rename_move(old_name, new_name)
        except (InvalidMoveName, DuplicateMoveError):
            return

    def global_remove_move(self, move_name):
        # self.detector_model.remove_move(move_name)
        self.controllers["image"].remove_move(move_name)
        self.controllers["phue"].remove_move(move_name)
        self.controllers["matrix"].remove_move(move_name)
        self.controllers["sk"].remove_move(move_name)
        self.controllers["detector"].remove_move(move_name)

    # ============================== Logic ============================== #
    def reset_models(self):
        for controller in self.controllers.values():
            controller.reset()

    def _initialize_app(self):
        # Set Project Path
        if not self.path_manager.set_folder_path():
            QMessageBox.critical(self.views["main"], "שגיאה", "לא נבחרה תיקייה")
            return False

        # Reset
        self.reset_models()

        # Initialize Controllers
        self.path_manager.set_files_path(self.models["phue"].phue_paths)
        self.controllers["settings"].init_model(self.path_manager.path_init)
        self.controllers["move"].init_model(self.path_manager.path_init_tk1)
        self.controllers["matrix"].init_model(self.path_manager.path_init_tk1)
        self.controllers["sk"].init_model(self.path_manager.path_init)
        self.controllers["io24"].init_model(self.path_manager.path_init)
        self.controllers["io64"].init_model(self.path_manager.path_init)
        self.controllers["detector"].init_model(self.path_manager.path_init_tk1)
        self.controllers["image"].init_model(self.path_manager.path_init_tk1, self.controllers["move"].get_all_moves_names())
        self.controllers["phue"].init_model(self.models["phue"].phue_paths, self.path_manager.path_init_tk1)
        self.controllers["schedule"].init_model(self.path_manager.path_init_tk1)
        self.controllers["parameters_ta"].init_model(self.path_manager.path_parameters_ta, len(self.models["image"].all_images))

        # self.path_manager.load_project_number_and_name()
        return True

    def start_new_project(self):
        reply = QMessageBox.question(
            self.views["main"],
            "אישור",
            "האם אתה בטוח?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No           # default
        )

        if reply == QMessageBox.StandardButton.No:
            return

        self.path_manager.path_project          = True
        self.views["schedule"].is_copy_sunday       = True
        self.views["navigator"].authority_combo.setDisabled(True)
        self.reset_models()
        self.show_view("settings")
        QMessageBox.information(self.views["main"],"פרויקט חדש","הופעל פרויקט חדש")

    def close_project(self):
        reply = QMessageBox.question(
            self.views["main"],
            "אישור",
            "האם אתה בטוח?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No           # default
        )

        if reply == QMessageBox.StandardButton.No:
            return

        self.show_view("")
        self.reset_models()
        self.path_manager.reset()
        self.path_manager.path_project          = None
        self.views["schedule"].is_copy_sunday       = True
        self.views["navigator"].authority_combo.setDisabled(False)

        QMessageBox.information(self.views["main"],"פרויקט חדש","הפרויקט נסגר")
        pass

    # ============================== Write to file ============================== #
    def write_to_code(self):
        if not self.is_data_valid():
            return

        if not self.path_manager.create_project(self.root):
            return

        # if maatz falcon


        # self.write_phase_imports()
        WriteDataManager.write_phase_imports(self.path_manager.path_init_tk1_dst, self.controllers["image"].get_all_images())

        self.controllers["settings"].write_settings_to_project(self.path_manager.path_init_dst)
        self.controllers["move"].write_moves_to_project(self.path_manager.path_tk1_dst, self.path_manager.path_init_tk1_dst)
        self.controllers["matrix"].write_matrix_to_file(self.path_manager.path_init_tk1_dst)
        self.controllers["sk"].write_to_file(self.path_manager.path_init_dst)
        self.controllers["schedule"].write_to_file(self.path_manager.path_init_tk1_dst)
        self.controllers["image"].write_to_file(self.path_manager.path_tk1_dst, self.path_manager.path_init_tk1_dst, self.path_manager.path_phase_folder_dst)
        self.controllers["phue"].write_to_file(self.path_manager.path_tk1_dst, self.path_manager.path_init_tk1_dst, self.path_manager.path_phue_folder_dst)
        self.controllers["parameters_ta"].write_to_file(self.path_manager.path_parameters_ta_dst, self.path_manager.path_init_tk1_dst, self.controllers["image"].fetch_images_by_sp())
        self.controllers["detector"].write_to_file(self.path_manager.path_init_tk1_dst, self.path_manager.path_tk1_dst)

        # פתיחת הקובץ
        doc = Document(self.path_manager.path_cards_dst)

        # תאריך נוכחי
        today = datetime.now().strftime("%d/%m/%Y")

        for section in doc.sections:
            header = section.header

            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if "<date>" in run.text:
                                    run.text = run.text.replace("<date>", today)        # שמירה

        doc.save("file_updated.docx")

    def is_data_valid(self):
        if not self.controllers["matrix"].is_matrix_valid():
            self.show_view("matrix")
            return False

        if not self.controllers["phue"].is_names_valid():
            self.show_view("phue")
            return False

        if not self.controllers["image"].is_sp_valid():
            self.show_view("image")
            return False

        if not self.controllers["detector"].is_data_valid():
            self.show_view("detector")
            return False
        return True

