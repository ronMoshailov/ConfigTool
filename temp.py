from docx import Document

doc = Document(r"C:\Users\ron.MENORAH-RND\Desktop\ta172\Docs\Cards.docx")
for i, section in enumerate(doc.sections):
    header = section.header
    print(f"--- Section {i} ---")

    print("PARAGRAPHS:")
    for p in header.paragraphs:
        print(p.text)

    print("TABLES:")
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)